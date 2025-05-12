"""
GUI tool to generate side-by-side screenshot comparisons of matched PDF pairs from two folders, export them as Word documents with proper layout, and convert to PDFs.
"""

import os
import fitz  # PyMuPDF
import PIL
from PIL import Image
import docx
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.shared import Pt
import docx2pdf
import time
from tkinter import font
from tkinter import messagebox
import tkinter as tk
from tkinter import filedialog

def create_folders_screenshots_and_PDF_file(source_path, target_path):    
    source_files = os.listdir(source_path)
    target_files = os.listdir(target_path)
    
    # Make sure the paths are different
    if source_path == target_path:
        show_error_same_paths()
        root.destroy()
    
    # Make sure both paths have the same number of files
    if len(source_files) != len(target_files):
        show_error_number_of_files()
        root.destroy()
    else:
    
        # Iterate over each pair of source and target files
        for source_file, target_file in zip(source_files, target_files):
            # Check if files are both directories
            if os.path.isdir(os.path.join(source_path, source_file)) or os.path.isdir(os.path.join(target_path, target_file)):
                continue

            # Check if files have the same name
            if os.path.splitext(source_file)[0] != os.path.splitext(target_file)[0]:
                show_error_same_name()
                root.destroy()
    
            folder_name = os.path.splitext(source_file)[0]
            folder_path = os.path.join(os.getcwd(), folder_name)
            os.makedirs(folder_path, exist_ok=True)
            
            source_doc_path = os.path.join(source_path, source_file)
            target_doc_path = os.path.join(target_path, target_file)
    
            # Take screenshots of all pages for both source and target documents
            take_screenshots_source(source_doc_path, folder_path, source_path, "source", source_files)
            page_sizes, source_whole_path, nozoom_folder, nozoom_images_whole_path = take_screenshots_source(source_doc_path, folder_path, source_path, "source", source_files)
            
            take_screenshots_target(target_doc_path, folder_path, target_path, "target", target_files)
            page_sizes, target_whole_path = take_screenshots_target(target_doc_path, folder_path, target_path, "target", target_files)
            
            
            all_image_path = []
            # Iterate over the zipped lists and alternate between elements
            for source, target in zip(source_whole_path, target_whole_path):
                all_image_path.append(source)
                all_image_path.append(target)
            
   
              
            doubled_page_sizes = []
            # Iterate over the page_sizes list
            for page_size in page_sizes:
                # Double the value and append it twice to the doubled_page_sizes list
                doubled_page_sizes.append(page_size)
                doubled_page_sizes.append(page_size)
            
            word_file_path = os.path.join(folder_path, f"{os.path.basename(folder_path)}.docx")
            
            # Create Word document with alternating screenshots
            create_alternating_word_doc(page_sizes, all_image_path, doubled_page_sizes, word_file_path, nozoom_images_whole_path)
                    
            # Remove existing PNG files
            remove_png_files(folder_path, nozoom_folder)    
            
            pdf_file_path = os.path.join(folder_name, source_file)
    
            # Convert to PDF
            convert_word_to_pdf(word_file_path, pdf_file_path)
    
            # Remove Word file
            remove_Word_file(folder_path)

def remove_Word_file(folder_path):
    for root, dirs, files in os.walk(folder_path):   
        for file in files:
            if file.endswith(".docx"):
                os.remove(os.path.join(root, file))

def convert_word_to_pdf(word_file_path, pdf_file_path):
    docx2pdf.convert(word_file_path, pdf_file_path)

  
def take_screenshots_source(source_doc_path, folder_path, source_path, prefix, source_files):  

    # Filter PDF files
    for x in range(0, 2):
        if x == 0:
            pdf_files = [file for file in source_files if file.endswith(".pdf")]
        else:
            pdf_files = [file for file in source_files if file.endswith(".PDF")]
    
        if pdf_files:
            # Process each PDF file found in the directory
            for pdf_file in pdf_files:
        
                pdf_file = fitz.open(source_doc_path)
                source_images = []
                nozoom_images = []
                page_sizes = []
                source_whole_path = []
                nozoom_images_whole_path = []
                
                # Iterate over pages and take screenshots
                for page_num in range(pdf_file.page_count):
                    page = pdf_file.load_page(page_num)
                        
                    nozoom_image = page.get_pixmap()
                                   
                    zoom = 2
                    mat = fitz.Matrix(zoom, zoom)
                    source_image = page.get_pixmap(matrix=mat)
                                   
                
                    nozoom_img = Image.frombytes("RGB", [nozoom_image.width, nozoom_image.height], nozoom_image.samples)
                    source_img = Image.frombytes("RGB", [source_image.width, source_image.height], source_image.samples)
                    
                    nozoom_folder = os.path.join(folder_path, 'nozoom_folder')
                    if not os.path.exists(nozoom_folder):
                        os.makedirs(nozoom_folder)
                    nozoom_image_path = f"nozoom_page_{page_num + 1:03d}_{prefix}.png"
                    source_image_path = f"page_{page_num + 1:03d}_{prefix}.png"
                    nozoom_img_path = f"{nozoom_folder}/{nozoom_image_path}"
                    source_img_path = f"{folder_path}/{source_image_path}"                                    
                                    
                    source_img.save(source_img_path, format="PNG")
                    nozoom_img.save(nozoom_img_path, format="PNG")
                    
                    source_whole_path.append(source_img_path)
                    source_images.append(source_image_path)
                    nozoom_images_whole_path.append(nozoom_img_path)
                    nozoom_images.append(nozoom_image_path)
                    page_sizes.append((page.rect.width, page.rect.height))
                    # page_size = page.rect.width, page.rect.height
                
                pdf_file.close()
        return page_sizes, source_whole_path, nozoom_folder, nozoom_images_whole_path
    
def take_screenshots_target(target_doc_path, folder_path, target_path, prefix, target_files):  

    # Filter PDF files
    for x in range(0, 2):
        if x == 0:
            pdf_files = [file for file in target_files if file.endswith(".pdf")]
        else:
            pdf_files = [file for file in target_files if file.endswith(".PDF")]    

        if pdf_files:
            # Process each PDF file found in the directory
            for pdf_file in pdf_files:
        
                pdf_file = fitz.open(target_doc_path)
                target_images = []
                page_sizes = []
                target_whole_path = []
                
                # Iterate over pages and take screenshots
                for page_num in range(pdf_file.page_count):
                    page = pdf_file.load_page(page_num)                       
                                                      
                    zoom = 2
                    mat = fitz.Matrix(zoom, zoom)
                    target_image = page.get_pixmap(matrix=mat)
                                   
            
                    target_img = Image.frombytes("RGB", [target_image.width, target_image.height], target_image.samples)
                    target_image_path = f"page_{page_num + 1:03d}_{prefix}.png"
                    
                    target_img_path = f"{folder_path}/{target_image_path}"                
                    target_img.save(target_img_path, format="PNG")
                    target_images.append(target_image_path)
                    
                    target_whole_path.append(target_img_path)
                    page_sizes.append((page.rect.width, page.rect.height))
                    # page_size = page.rect.width, page.rect.height
                
                pdf_file.close()
            
        return page_sizes, target_whole_path
        
def create_alternating_word_doc(page_sizes, all_image_path, doubled_page_sizes, word_file_path, nozoom_images_whole_path):

    # Determine the layout for the first image in the images list
    first_page_landscape = check_first_image_layout(page_sizes)
    
    # first_image = check_first_image_layout(page_sizes) 
    
    if first_page_landscape == False:
        create_word_document_first_portrait(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path, page_sizes, nozoom_images_whole_path)
        return first_page_landscape
    else:
        create_word_document_first_landscape(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path)
        return first_page_landscape
    

def check_first_image_layout(page_sizes):
     if page_sizes:
         first_image = page_sizes[0]
         width, height = first_image
         return width > height  # True if landscape, False if portrait
     else:
         return None  # Return None for an empty image list 
     return first_image

def create_word_document_first_portrait(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path, page_sizes, nozoom_images_whole_path):
    doc = Document()
    sections = doc.sections
    landscape_page = first_page_landscape
    # first_image = page_sizes[0]
    first_image = PIL.Image.open(nozoom_images_whole_path[0])
    first_image_width, first_image_height = first_image.size
    
    # Set margins to zero (remove margins)
    for section in sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0)

    section.page_width = Pt(first_image_width)
    section.page_height = Pt(first_image_height)

    for image, page_size in zip(all_image_path, doubled_page_sizes):
        width, height = page_size
            
        if width > height:  # Check if width > height (indicating landscape layout)
            if landscape_page == False:
                section = doc.add_section(WD_ORIENT.LANDSCAPE)
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            landscape_page = True    
        else:
            if landscape_page == True:
                section = doc.add_section(WD_ORIENT.PORTRAIT)
                section.page_width = Pt(width)
                section.page_height = Pt(height)
            landscape_page = False

           
        doc.add_picture(image, width=Pt(width)) # Set picture width based on image width

    doc.save(word_file_path)    
            
        
def create_word_document_first_landscape(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path):
    doc = Document()
    sections = doc.sections
    section = doc.sections[-1]
    landscape_page = first_page_landscape
    image_number = 1
    

    for image, page_size in zip(all_image_path, doubled_page_sizes):
        width, height = page_size  

        if image_number == 1:
            section.page_width = Pt(width)
            section.page_height = Pt(height)
       
        if width > height:  # Check if width > height (indicating landscape layout)
            if landscape_page == False:
                section = doc.add_section(WD_ORIENT.LANDSCAPE)
                section.page_width = Pt(width)
                section.page_height = Pt(height)
            landscape_page = True    
        else:
            if landscape_page == True:
                section = doc.add_section(WD_ORIENT.PORTRAIT)
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            landscape_page = False
        
        doc.add_picture(image, width=Pt(width)) # Set picture width based on image width

        image_number += 1

    # Set margins to zero (remove margins)
    for section in sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0) 
        
    doc.save(word_file_path)
    

def remove_png_files(folder_path, nozoom_folder):
    for root, dirs, files in os.walk(folder_path, topdown=False):      
        for subdir in dirs:
                    subdir_path = nozoom_folder
                    if not os.listdir(subdir_path):
                        os.rmdir(subdir_path)
        for file in files:
            if file.endswith(".png"):
                os.remove(os.path.join(root, file))

def get_source_path():
    source_path = filedialog.askdirectory()
    source_path_entry.delete(0, tk.END)
    source_path_entry.insert(0, source_path)

def get_target_path():
    target_path = filedialog.askdirectory()
    target_path_entry.delete(0, tk.END)
    target_path_entry.insert(0, target_path)
    

def show_error_same_paths():
       
    messagebox.showerror("Error", "Error: The source and target paths must be different.")
    root.destroy() 
    
    
def show_error_number_of_files():
    
    messagebox.showerror("Error", "Error: The number of files in the source and target paths\ndoes not match.")
    root.destroy()
 
def show_error_same_name():
    
    messagebox.showerror("Error", "Error: The source and target files must have the same name.")
    root.destroy()
    
        
# Function to show "Finished" message with elapsed time
def show_finished_message(start_time):
    elapsed_time = time.time() - start_time
    minutes = int(elapsed_time // 60)
    seconds = int(elapsed_time % 60)
    messagebox.showinfo("Script Finished", f"All files have been processed.\nScript run time: {minutes:02d}:{seconds:02d}")

def process_paths():
       
    start_time = time.time()  # Start the stopwatch
   
    source_path = source_path_entry.get()
    target_path = target_path_entry.get()
    create_folders_screenshots_and_PDF_file(source_path, target_path)
   
    show_finished_message(start_time)  # Show "Finished" message with elapsed time   

    root.destroy()


# Create tkinter window
root = tk.Tk()
root.withdraw() 
root.title("Document Screenshot Tool")
width = 478
height = 72
screenwidth = root.winfo_screenwidth()
screenheight = root.winfo_screenheight()
root.geometry()

root.eval('tk::PlaceWindow . center')
font_normal = font.Font(family = 'Arial', size = 10)

# Source path entry
source_path_label = tk.Label(root, text="Source Path:")
source_path_label.grid(row=0, column=0, padx=5, pady=5)
source_path_entry = tk.Entry(root, width=50)
source_path_entry.grid(row=0, column=1, padx=5, pady=5)
source_path_button = tk.Button(root, text="Browse", command=get_source_path)
source_path_button.grid(row=0, column=2, padx=5, pady=5)

# Target path entry
target_path_label = tk.Label(root, text="Target Path:")
target_path_label.grid(row=1, column=0, padx=5, pady=5)
target_path_entry = tk.Entry(root, width=50)
target_path_entry.grid(row=1, column=1, padx=5, pady=5)
target_path_button = tk.Button(root, text="Browse", command=get_target_path)
target_path_button.grid(row=1, column=2, padx=5, pady=5)

# Process button
process_button = tk.Button(root, text="Process", command=process_paths)
process_button.grid(row=2, column=1, padx=5, pady=10)

root.mainloop()
