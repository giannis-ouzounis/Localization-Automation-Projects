"""
Rendering each PDF page as a PNG image.
Embedding those images in a Word .docx file.
Adjusting the page orientation (portrait/landscape) based on image dimensions.
Offering a minimal GUI for folder input and execution via Tkinter.
"""
import os
import fitz  # PyMuPDF
from PIL import Image
import PIL
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.shared import Pt
from tkinter import *
from tkinter import font
from tkinter import messagebox

def check_first_image_layout(images):
    if images:
        first_image = PIL.Image.open(images[0])
        width, height = first_image.size
        return width > height  # True if landscape, False if portrait
    else:
        return None  # Return None for an empty image list

def create_word_document_first_portrait(images, page_sizes, first_page_landscape, pdf_name):
    doc = Document()
    sections = doc.sections
    landscape_page = first_page_landscape
    first_image = PIL.Image.open(images[0])
    first_image_width, first_image_height = first_image.size

    # Set margins to zero (remove margins)
    for section in sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0)

    section.page_width = Pt(first_image_width)
    section.page_height = Pt(first_image_height)

    for img_path, size in zip(images, page_sizes):
        img = PIL.Image.open(img_path)
        width, height = img.size
        
        if width > height:  # Check if width > height (indicating landscape layout)
            if landscape_page == False:
                section = doc.add_section(WD_ORIENT.LANDSCAPE)
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            landscape_page = True    
        else:
            if landscape_page == True:
                section = doc.add_section(WD_ORIENT.PORTRAIT)
                section.page_width = Pt(width)
                section.page_height = Pt(height)
            landscape_page = False
        
        doc.add_picture(img_path, width=Pt(width))  # Set picture width based on image width
              
    doc.save(pdf_name + '.docx')
    
def create_word_document_first_landscape(images, page_sizes, first_page_landscape, pdf_name):
    doc = Document()
    sections = doc.sections
    section = doc.sections[-1]
    landscape_page = first_page_landscape
    image_number = 1
    
    for img_path, size in zip(images, page_sizes):
        img = PIL.Image.open(img_path)
        width, height = img.size  

        if image_number == 1:
            section.page_width = Pt(width)
            section.page_height = Pt(height)
       
        if width > height:  # Check if width > height (indicating landscape layout)
            if landscape_page == False:
                section = doc.add_section(WD_ORIENT.LANDSCAPE)
                section.page_width = Pt(width)
                section.page_height = Pt(height)
            landscape_page = True    
        else:
            if landscape_page == True:
                section = doc.add_section(WD_ORIENT.PORTRAIT)
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            landscape_page = False
        
        doc.add_picture(img_path, width=Pt(width))  # Set picture width based on image width
    
        image_number += 1        
    
    # Set margins to zero (remove margins)
    for section in sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0)
    
    doc.save(pdf_name + '.docx')

def remove_png_files(current_script_path):
    png_files = [file for file in os.listdir(current_script_path) if file.endswith(".png")]

    for png_file in png_files:
        file_path = os.path.join(current_script_path, png_file)
        os.remove(file_path)

def Start():
    # Get the current working directory
    current_directory = entry_path.get()
    os.chdir(current_directory)
    
    # List all files in the current directory
    files_in_directory = os.listdir(current_directory)
    
    # Filter PDF files
    for x in range(0, 2):
        if x == 0:
            pdf_files = [file for file in files_in_directory if file.endswith(".pdf")]
        else:
            pdf_files = [file for file in files_in_directory if file.endswith(".PDF")]
    
        if pdf_files:
            # Process each PDF file found in the directory
            for pdf_file in pdf_files:
                pdf_name = os.path.splitext(pdf_file)[0]  # Extracting PDF name without extension
                pdf_path = os.path.join(current_directory, pdf_file)
                pdf_document = fitz.open(pdf_path)
                images = []
                page_sizes = []
                
                for page_num in range(pdf_document.page_count):
                    page = pdf_document.load_page(page_num)
                    image = page.get_pixmap()
                    img = PIL.Image.frombytes("RGB", [image.width, image.height], image.samples)
                    image_path = f"page_{page_num + 1}.png"
                    img.save(image_path, "PNG")
                    images.append(image_path)
                    page_sizes.append((page.rect.width, page.rect.height))
                
                pdf_document.close()
                
                # Determine the layout for the first image in the images list
                first_page_landscape = check_first_image_layout(images)
                
                if first_page_landscape == False:
                    create_word_document_first_portrait(images, page_sizes, first_page_landscape, pdf_name)
                else:
                    create_word_document_first_landscape(images, page_sizes, first_page_landscape, pdf_name)
                
                # current_script_path = os.path.dirname(__file__)
                
                remove_png_files(entry_path.get())
    messagebox.showinfo('', 'Finished.')

window = Tk() # Add UI for this script
width = 478
height = 72
screenwidth = window.winfo_screenwidth()
screenheight = window.winfo_screenheight()
window.title('PDF into Word screenshots')
window.resizable(False, False)
window.geometry('%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2))

font_normal = font.Font(family = 'Arial', size = 10)

label_path = Label(window, text = 'Please specify a folder that contains all PDF files', font = font_normal)
label_path.place(x = 10, y = 10)
entry_path = Entry(window, width = 50, font = font_normal)
entry_path.place(x = 12, y = 35)
button_start = Button(window, text = 'Start', width = 10, font = font_normal,command = Start)
button_start.place(x = 378, y = 31)

window.mainloop()