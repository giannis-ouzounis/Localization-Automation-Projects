"""
Allows interactive mapping of multiple source PDF files to one or more target PDF files using a Tkinter GUI.
Supports caching via MD5 checksums to skip redundant screenshot generation.
Converts each page of source and target PDFs to high-resolution images using PyMuPDF (with zoom applied).
Generates side-by-side page comparisons, alternating between source and target pages.
Builds a Word document with automatic landscape/portrait layout detection for each page.
Converts the Word document into PDF output using docx2pdf.
Outputs results into structured subfolders, one per target, for easy review.
Includes scrollable mapping display for tracking selected file pairings.
Final cleanup: deletes temporary PNG and DOCX files, keeping only final PDFs.
"""

import os
import fitz  #PyMuPDF
import PIL
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.shared import Pt
import docx2pdf
import time
from tkinter import font
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import shutil
import hashlib

class FileMapperApp:
    def __init__(self, root):
        self.root = root
        self.root.title("File Mapper")
        self.root.minsize(800, 400)  # Set the minimum size of the window

        self.source_files = []
        self.target_files = []
        self.selected_source_index = None
        self.selected_target_indices = []
        self.mappings = []

        self.create_widgets()

    def create_widgets(self):
        
        frame = ttk.Frame(self.root, padding="10")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Source Listbox and Scrollbars
        ttk.Label(frame, text="Source Files").grid(row=0, column=0, padx=5, pady=5)

        source_frame = ttk.Frame(frame)
        source_frame.grid(row=1, column=0, rowspan=2, padx=5, pady=5, sticky='nsew')

        self.source_listbox = tk.Listbox(source_frame, selectmode=tk.EXTENDED, width=80)
        self.source_listbox.grid(row=0, column=0, sticky='nsew')
        self.source_listbox.bind('<<ListboxSelect>>', self.update_selected_source)

        source_vsb = ttk.Scrollbar(source_frame, orient="vertical", command=self.source_listbox.yview)
        source_vsb.grid(row=0, column=1, sticky='ns')
        self.source_listbox.configure(yscrollcommand=source_vsb.set)

        source_hsb = ttk.Scrollbar(source_frame, orient="horizontal", command=self.source_listbox.xview)
        source_hsb.grid(row=1, column=0, sticky='ew')
        self.source_listbox.configure(xscrollcommand=source_hsb.set)

        ttk.Button(frame, text="Add Source Files", command=self.add_source_files).grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(frame, text="Remove Source Files", command=self.remove_source_files).grid(row=2, column=1, padx=5, pady=5)

        # Target Listbox and Scrollbars
        ttk.Label(frame, text="Target Files").grid(row=3, column=0, padx=5, pady=5)

        target_frame = ttk.Frame(frame)
        target_frame.grid(row=4, column=0, rowspan=2, padx=5, pady=5, sticky='nsew')

        self.target_listbox = tk.Listbox(target_frame, selectmode=tk.EXTENDED, width=80)
        self.target_listbox.grid(row=0, column=0, sticky='nsew')
        self.target_listbox.bind('<<ListboxSelect>>', self.update_selected_targets)

        target_vsb = ttk.Scrollbar(target_frame, orient="vertical", command=self.target_listbox.yview)
        target_vsb.grid(row=0, column=1, sticky='ns')
        self.target_listbox.configure(yscrollcommand=target_vsb.set)

        target_hsb = ttk.Scrollbar(target_frame, orient="horizontal", command=self.target_listbox.xview)
        target_hsb.grid(row=1, column=0, sticky='ew')
        self.target_listbox.configure(xscrollcommand=target_hsb.set)

        ttk.Button(frame, text="Add Target Files", command=self.add_target_files).grid(row=4, column=1, padx=5, pady=5)
        ttk.Button(frame, text="Remove Target Files", command=self.remove_target_files).grid(row=5, column=1, padx=5, pady=5)

        ttk.Button(frame, text="Map Files", command=self.map_files).grid(row=6, column=0, columnspan=2, pady=5)

        # Mapping display and Scrollbars
        mapping_frame = ttk.Frame(frame)
        mapping_frame.grid(row=7, column=0, columnspan=2, pady=5, sticky='nsew')

        self.mapping_display = tk.Text(mapping_frame, width=100, height=10)
        self.mapping_display.grid(row=0, column=0, sticky='nsew')

        mapping_vsb = ttk.Scrollbar(mapping_frame, orient="vertical", command=self.mapping_display.yview)
        mapping_vsb.grid(row=0, column=1, sticky='ns')
        self.mapping_display.configure(yscrollcommand=mapping_vsb.set)

        mapping_hsb = ttk.Scrollbar(mapping_frame, orient="horizontal", command=self.mapping_display.xview)
        mapping_hsb.grid(row=1, column=0, sticky='ew')
        self.mapping_display.configure(xscrollcommand=mapping_hsb.set)

        ttk.Button(frame, text="Remove Selected Mapping", command=self.remove_selected_mapping).grid(row=8, column=0, columnspan=2, pady=5)

        ttk.Button(frame, text="OK", command=process_mappings).grid(row=9, column=0, columnspan=2, pady=5)

    def add_source_files(self):
        files = filedialog.askopenfilenames(title="Select Source Files")
        for file in files:
            self.source_files.append(file)
            self.source_listbox.insert(tk.END, file)

    def add_target_files(self):
        files = filedialog.askopenfilenames(title="Select Target Files")
        for file in files:
            self.target_files.append(file)
            self.target_listbox.insert(tk.END, file)

    def remove_source_files(self):
        selected_indices = self.source_listbox.curselection()
        for index in selected_indices[::-1]:  # Reverse the order to avoid issues while deleting
            self.source_listbox.delete(index)
            del self.source_files[index]

    def remove_target_files(self):
        selected_indices = self.target_listbox.curselection()
        for index in selected_indices[::-1]:  # Reverse the order to avoid issues while deleting
            self.target_listbox.delete(index)
            del self.target_files[index]

    def update_selected_source(self, event):
        selection = self.source_listbox.curselection()
        if selection:
            self.selected_source_index = selection[0]

    def update_selected_targets(self, event):
        self.selected_target_indices = self.target_listbox.curselection()

    def map_files(self):
        if self.selected_source_index is None:
            messagebox.showerror("Error", "Please select a source file.")
            return

        if not self.selected_target_indices:
            messagebox.showerror("Error", "Please select at least one target file.")
            return

        source_file = self.source_files[self.selected_source_index]
        target_files = [self.target_files[i] for i in self.selected_target_indices]

        mapping = f"{source_file} -> {', '.join(target_files)}\n"
        self.mapping_display.insert(tk.END, mapping)
        self.mappings.append((source_file, target_files))
        return source_file, target_files

    def remove_selected_mapping(self):
        try:
            selected_text = self.mapping_display.selection_get()
            if selected_text:
                start_index = self.mapping_display.search(selected_text, "1.0", tk.END)
                end_index = f"{start_index} + {len(selected_text)}c"
                self.mapping_display.delete(start_index, end_index)
                
                # Also remove from mappings list
                source_file, target_files = selected_text.strip().split(" -> ")
                target_files = target_files.split(", ")
                self.mappings.remove((source_file, target_files))
                
        except tk.TclError:
            messagebox.showerror("Error", "Please select a mapping to remove.")
            
    def get_mappings(self):
        return self.mappings

# Dictionary to store cached results
cache = {}

def compute_md5(file_path):
    """Compute the MD5 hash of a file."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def create_folders_screenshots_and_PDF_file(mappings):     
            
    # Iterate over each pair of source and target files
    for mapping in mappings:
        source_path, target_paths = mapping

        # source_folder is the path to the folder with all source files
        source_folder = os.path.dirname(source_path)
        
        # os.path.basename returns the filename with the extension
        source_file = os.path.basename(source_path)
        source_files = os.listdir(source_folder)
        
        # os.path.splitext returns the filename without the extension
        source_name = os.path.splitext(source_file)[0]
        
        # source_folder_path is the newly created folder for each of the source files inside the source folder
        source_folder_path = os.path.join(source_folder, source_name)
        os.makedirs(source_folder_path, exist_ok=True)
                
        # Take screenshots of all pages for source documents
        page_sizes, source_whole_path, nozoom_folder, nozoom_images_whole_path = take_screenshots_source(source_folder_path, source_path, "source", source_files)

        
        for target_path in target_paths:
            target_folder = os.path.dirname(target_path)
            target_files = os.listdir(target_folder)
            target_file = os.path.basename(target_path)
            
            # os.path.splitext returns the path and the filename without the extension
            target_name = os.path.splitext(target_file)[0]
            # folder_name is the name of each of the target files without the extension
            folder_name = target_name
            # folder_path is the path of each of the folders were the output file will be saved, same directory as the script is
            folder_path = os.path.join(os.getcwd(), folder_name)

            os.makedirs(folder_path, exist_ok=True)
          
             
            # Take screenshots of all pages for target documents
            page_sizes, target_whole_path = take_screenshots_target(folder_path, target_path, "target", target_files)
            
            all_image_path = [img for pair in zip(source_whole_path, target_whole_path) for img in pair]                     
            
            doubled_page_sizes = [size for page_size in page_sizes for size in [page_size, page_size]]
            
            word_file_path = os.path.join(folder_path, f"{os.path.basename(folder_path)}.docx")
            
            # Create Word document with alternating screenshots
            create_alternating_word_doc(page_sizes, all_image_path, doubled_page_sizes, word_file_path, nozoom_images_whole_path)
                    
            # Remove existing PNG files
            remove_png_files(folder_path, nozoom_folder)    
            
            pdf_file_path = os.path.join(folder_name, source_file)
    
            # Convert to PDF
            convert_word_to_pdf(word_file_path, pdf_file_path)
    
            # Remove Word file
            remove_Word_file(folder_path)
            
        # Remove existing PNG files
        remove_source_png_files(source_folder_path)

def remove_Word_file(folder_path):
    for root, dirs, files in os.walk(folder_path):   
        for file in files:
            if file.endswith(".docx"):
                os.remove(os.path.join(root, file))

def convert_word_to_pdf(word_file_path, pdf_file_path):
    docx2pdf.convert(word_file_path, pdf_file_path)

  
def take_screenshots_source(source_folder_path, source_path, prefix, source_files):  

    # Filter PDF files
    for x in range(2):
        pdf_files = [file for file in source_files if file.endswith(".pdf") or file.endswith(".PDF")]
    
        if pdf_files:
            # Process each PDF file found in the directory
            for pdf_file in pdf_files:
                
                # Compute the hash of the PDF file
                pdf_hash = compute_md5(source_path)
                
                # Check if the result is already in the cache
                if pdf_hash in cache:                    
                    return cache[pdf_hash]
                
                # Otherwise, perform the screenshot operation
        
                pdf_file = fitz.open(source_path)
                source_images = []
                nozoom_images = []
                page_sizes = []
                source_whole_path = []
                nozoom_images_whole_path = []
                
                # Iterate over pages and take screenshots
                for page_num in range(pdf_file.page_count):
                    page = pdf_file.load_page(page_num)
                        
                    nozoom_image = page.get_pixmap()
                                   
                    zoom = 2
                    mat = fitz.Matrix(zoom, zoom)
                    source_image = page.get_pixmap(matrix=mat)
                                   
                
                    nozoom_img = Image.frombytes("RGB", [nozoom_image.width, nozoom_image.height], nozoom_image.samples)
                    source_img = Image.frombytes("RGB", [source_image.width, source_image.height], source_image.samples)
                    
                    nozoom_folder = os.path.join(source_folder_path, 'nozoom_folder')
                    if not os.path.exists(nozoom_folder):
                        os.makedirs(nozoom_folder)
                    nozoom_image_path = f"nozoom_page_{page_num + 1:03d}_{prefix}.png"
                    source_image_path = f"page_{page_num + 1:03d}_{prefix}.png"
                    nozoom_img_path = f"{nozoom_folder}/{nozoom_image_path}"
                    source_img_path = f"{source_folder_path}/{source_image_path}"                                    
                                    
                    source_img.save(source_img_path, format="PNG")
                    nozoom_img.save(nozoom_img_path, format="PNG")
                    
                    source_whole_path.append(source_img_path)
                    source_images.append(source_image_path)
                    nozoom_images_whole_path.append(nozoom_img_path)
                    nozoom_images.append(nozoom_image_path)
                    page_sizes.append((page.rect.width, page.rect.height))
                
                pdf_file.close()
                
                # Store the result in the cache
                cache[pdf_hash] = (page_sizes, source_whole_path, nozoom_folder, nozoom_images_whole_path)
                
            return page_sizes, source_whole_path, nozoom_folder, nozoom_images_whole_path
    
def take_screenshots_target(folder_path, target_path, prefix, target_files):  

    # Filter PDF files         
    for x in range(2):
        pdf_files = [file for file in target_files if file.endswith(".pdf") or file.endswith(".PDF")]

        if pdf_files:
            # Process each PDF file found in the directory
            for pdf_file in pdf_files:
                
                # Compute the hash of the PDF file
                pdf_hash = compute_md5(target_path)
                
                # Check if the result is already in the cache
                if pdf_hash in cache:                    
                    return cache[pdf_hash]
                
                # Otherwise, perform the screenshot operation
                        
                pdf_file = fitz.open(target_path)
                target_images = []
                page_sizes = []
                target_whole_path = []
                
                # Iterate over pages and take screenshots
                for page_num in range(pdf_file.page_count):
                    page = pdf_file.load_page(page_num)                       
                                                      
                    zoom = 2
                    mat = fitz.Matrix(zoom, zoom)
                    target_image = page.get_pixmap(matrix=mat)
                                   
            
                    target_img = Image.frombytes("RGB", [target_image.width, target_image.height], target_image.samples)
                    target_image_path = f"page_{page_num + 1:03d}_{prefix}.png"
                    
                    target_img_path = f"{folder_path}/{target_image_path}"                
                    target_img.save(target_img_path, format="PNG")
                    target_images.append(target_image_path)
                    
                    target_whole_path.append(target_img_path)
                    page_sizes.append((page.rect.width, page.rect.height))
                    
                    # Store the result in the cache
                    cache[pdf_hash] = (page_sizes, target_whole_path)
                
                pdf_file.close()
            
        return page_sizes, target_whole_path
        
def create_alternating_word_doc(page_sizes, all_image_path, doubled_page_sizes, word_file_path, nozoom_images_whole_path):

    # Determine the layout for the first image in the images list
    first_page_landscape = check_first_image_layout(page_sizes)
    
    # first_image = check_first_image_layout(page_sizes) 
    
    if first_page_landscape == False:
        create_word_document_first_portrait(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path, page_sizes, nozoom_images_whole_path)
        return first_page_landscape
    else:
        create_word_document_first_landscape(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path)
        return first_page_landscape
    
def check_first_image_layout(page_sizes):
     if page_sizes:
         first_image = page_sizes[0]
         width, height = first_image
         return width > height  # True if landscape, False if portrait
     else:
         return None  # Return None for an empty image list 
     return first_image

def create_word_document_first_portrait(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path, page_sizes, nozoom_images_whole_path):
    doc = Document()
    sections = doc.sections
    landscape_page = first_page_landscape
    # first_image = page_sizes[0]
    first_image = PIL.Image.open(nozoom_images_whole_path[0])
    first_image_width, first_image_height = first_image.size
    
    # Set margins to zero (remove margins)
    for section in sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0)

    section.page_width = Pt(first_image_width)
    section.page_height = Pt(first_image_height)

    for image, page_size in zip(all_image_path, doubled_page_sizes):
        width, height = page_size
            
        if width > height:  # Check if width > height (indicating landscape layout)
            if landscape_page == False:
                section = doc.add_section(WD_ORIENT.LANDSCAPE)
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            landscape_page = True    
        else:
            if landscape_page == True:
                section = doc.add_section(WD_ORIENT.PORTRAIT)
                section.page_width = Pt(width)
                section.page_height = Pt(height)
            landscape_page = False

           
        doc.add_picture(image, width=Pt(width)) # Set picture width based on image width

    doc.save(word_file_path)    
            
        
def create_word_document_first_landscape(first_page_landscape, all_image_path, doubled_page_sizes, word_file_path):
    doc = Document()
    sections = doc.sections
    section = doc.sections[-1]
    landscape_page = first_page_landscape
    image_number = 1
    

    for image, page_size in zip(all_image_path, doubled_page_sizes):
        width, height = page_size  

        if image_number == 1:
            section.page_width = Pt(width)
            section.page_height = Pt(height)
       
        if width > height:  # Check if width > height (indicating landscape layout)
            if landscape_page == False:
                section = doc.add_section(WD_ORIENT.LANDSCAPE)
                section.page_width = Pt(width)
                section.page_height = Pt(height)
            landscape_page = True    
        else:
            if landscape_page == True:
                section = doc.add_section(WD_ORIENT.PORTRAIT)
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            landscape_page = False
        
        doc.add_picture(image, width=Pt(width)) # Set picture width based on image width

        image_number += 1

    # Set margins to zero (remove margins)
    for section in sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0) 
        
    doc.save(word_file_path)
    

def remove_png_files(folder_path, nozoom_folder):
    for root, dirs, files in os.walk(folder_path, topdown=False):      
        for subdir in dirs:
                    subdir_path = nozoom_folder
                    if not os.listdir(subdir_path):
                        os.rmdir(subdir_path)
        for file in files:
            if file.endswith(".png"):
                os.remove(os.path.join(root, file))


def remove_source_png_files(source_folder_path):              
    # Check if the directory exists and remove it
    if os.path.isdir(source_folder_path):
        shutil.rmtree(source_folder_path)  # Removes the directory and all its contents
    
        
# Function to show "Finished" message with elapsed time
def show_finished_message(start_time):
    elapsed_time = time.time() - start_time
    minutes = int(elapsed_time // 60)
    seconds = int(elapsed_time % 60)
    messagebox.showinfo("Script Finished", f"All files have been processed.\nScript run time: {minutes:02d}:{seconds:02d}")

def process_mappings():
       
    start_time = time.time()  # Start the stopwatch
    mappings = app.get_mappings()

    create_folders_screenshots_and_PDF_file(mappings)
   
    show_finished_message(start_time)  # Show "Finished" message with elapsed time   

    root.destroy()


# Create tkinter window
root = tk.Tk()
app = FileMapperApp(root)

root.withdraw() 
root.title("Document Screenshot Tool")
width = 478
height = 72
screenwidth = root.winfo_screenwidth()
screenheight = root.winfo_screenheight()
root.geometry()

root.eval('tk::PlaceWindow . center')
font_normal = font.Font(family = 'Arial', size = 10)

root.mainloop()
